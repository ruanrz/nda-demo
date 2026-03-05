# -*- coding: utf-8 -*-
"""
Word document generator with native Word Track Changes.

Produces a .docx using Word revision elements:
  - Deleted text → <w:del> (real deletion revision)
  - Added text   → <w:ins> (real insertion revision)
  - Unchanged    → normal runs

Uses diff_match_patch for word-level diffing: each unique word/whitespace
token is mapped to a single Unicode character so that diff_match_patch
operates on whole words rather than individual characters.
When rationale cannot be encoded in revision metadata, attach one concise
comment per changed paragraph/revision anchor.
"""

import copy
import io
import re
from datetime import datetime, timezone
from difflib import SequenceMatcher
from typing import List, Dict, Any, Optional

import diff_match_patch as dmp_module
from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


# ── Colours ──────────────────────────────────────────────────────
_RED = RGBColor(0xCF, 0x22, 0x2E)
_GREEN = RGBColor(0x1A, 0x7F, 0x37)
_BLACK = RGBColor(0x00, 0x00, 0x00)
_GREY = RGBColor(0x66, 0x66, 0x66)

_PARA_PUNCT_MAP = str.maketrans({
    "\u201c": '"',
    "\u201d": '"',
    "\u2018": "'",
    "\u2019": "'",
    "\u2014": "-",
    "\u2013": "-",
    "\u00a0": " ",
})

_WORD_TOKEN_RE = re.compile(r"\w+|[^\w\s]|\s+")


def generate_redline_docx(
    original_text: str,
    modified_text: str,
    issues_list: Optional[List[Dict[str, Any]]] = None,
    modifications: Optional[List[Dict[str, Any]]] = None,
    revisions: Optional[List[Dict[str, Any]]] = None,
    insertions: Optional[List[Dict[str, Any]]] = None,
    title: Optional[str] = None,
    redline_heading: Optional[str] = None,
    include_issues_list: bool = False,
    source_docx_bytes: Optional[bytes] = None,
) -> io.BytesIO:
    """
    Generate a Word document with native Track Changes markup.

    Args:
        original_text:  The original contract text.
        modified_text:  The modified contract text.
        issues_list:    Optional issues list to append as a table.
        modifications:  Optional modification list used for short rationale comments.
        title:          Optional document title heading (H1). If None, no title heading.
        redline_heading: Optional section heading before redlined body (H2). If None, no section heading.
        include_issues_list: Whether to append Issues List table at the end.
        source_docx_bytes: Optional original .docx bytes used as template to preserve
                          paragraph-level styles (including native numbering).

    Returns:
        BytesIO buffer containing the .docx file.
    """
    use_source_template = False
    if source_docx_bytes:
        try:
            doc = Document(io.BytesIO(source_docx_bytes))
            use_source_template = True
        except Exception:
            doc = Document()
    else:
        doc = Document()

    if not use_source_template:
        # ── Styles ───────────────────────────────────────────────
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        # ── Optional headings ───────────────────────────────────
        # Keep disabled by default so exported redline mirrors contract structure.
        if title:
            doc.add_heading(title, level=1)
        if redline_heading:
            doc.add_heading(redline_heading, level=2)

        # ── Redline body (plain mode) ───────────────────────────
        _add_redline_paragraphs(doc, original_text, modified_text, modifications or [], revisions=revisions)
    else:
        # ── Redline body (template mode; keeps paragraph properties) ──
        _add_redline_paragraphs_from_template(
            doc, original_text, modified_text, modifications or [], revisions=revisions
        )

    # ── Inserted clauses (Step 2b) — added as pure tracked insertions ──
    if insertions:
        _add_insertion_paragraphs(doc, insertions)

    # ── Issues List table ────────────────────────────────────────
    if include_issues_list and issues_list:
        doc.add_page_break()
        doc.add_heading("Issues List", level=2)
        _add_issues_table(doc, issues_list)

    # Ensure Word opens this file in a review-friendly state so native
    # w:ins / w:del revisions are visible by default.
    _ensure_review_settings(doc)

    # ── Write to buffer ──────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def compute_word_diffs(text1: str, text2: str):
    """Word-level diff using diff_match_patch.

    Tokenizes both texts into words, punctuation, and whitespace runs,
    maps each unique token to a single Unicode character, diffs the encoded
    strings, then decodes back.  The result has the same (op, text) format
    as ``dmp.diff_main`` but every chunk is aligned to token boundaries.

    A refinement pass extracts common token-level prefixes/suffixes from
    adjacent DELETE/INSERT pairs so that unchanged words are never marked
    as edited.
    """
    dmp = dmp_module.diff_match_patch()
    dmp.Diff_Timeout = 120

    tokens1 = _WORD_TOKEN_RE.findall(text1) if text1 else []
    tokens2 = _WORD_TOKEN_RE.findall(text2) if text2 else []
    if not tokens1 and not tokens2:
        return []

    token_array: List[str] = [""]
    token_map: Dict[str, str] = {}

    def _encode(tokens: List[str]) -> str:
        chars: List[str] = []
        for tok in tokens:
            if tok not in token_map:
                token_array.append(tok)
                token_map[tok] = chr(len(token_array) - 1)
            chars.append(token_map[tok])
        return "".join(chars)

    enc1 = _encode(tokens1)
    enc2 = _encode(tokens2)

    diffs = dmp.diff_main(enc1, enc2)
    dmp.diff_cleanupSemantic(diffs)

    result = []
    for op, data in diffs:
        decoded = "".join(token_array[ord(c)] for c in data)
        result.append((op, decoded))

    result = _refine_diffs(result)
    result = _merge_adjacent_diffs(result)
    return result


def _refine_diffs(
    diffs: List[tuple],
) -> List[tuple]:
    """Extract common token-level prefixes/suffixes from DELETE/INSERT pairs.

    When diff_match_patch marks "Company" as deleted and "Company, as defined
    by applicable law," as inserted, this function detects that "Company" is
    a shared prefix and splits the pair into EQUAL("Company") + INSERT(rest),
    producing minimal, precise redlines.
    """
    EQUAL = dmp_module.diff_match_patch.DIFF_EQUAL
    DELETE = dmp_module.diff_match_patch.DIFF_DELETE
    INSERT = dmp_module.diff_match_patch.DIFF_INSERT

    refined: List[tuple] = []
    i = 0
    while i < len(diffs):
        if (
            i + 1 < len(diffs)
            and diffs[i][0] == DELETE
            and diffs[i + 1][0] == INSERT
        ):
            del_text = diffs[i][1]
            ins_text = diffs[i + 1][1]

            del_tokens = _WORD_TOKEN_RE.findall(del_text)
            ins_tokens = _WORD_TOKEN_RE.findall(ins_text)

            prefix_len = 0
            for a, b in zip(del_tokens, ins_tokens):
                if a == b:
                    prefix_len += 1
                else:
                    break

            del_rest = del_tokens[prefix_len:]
            ins_rest = ins_tokens[prefix_len:]
            suffix_len = 0
            for a, b in zip(reversed(del_rest), reversed(ins_rest)):
                if a == b:
                    suffix_len += 1
                else:
                    break

            if prefix_len > 0:
                refined.append((EQUAL, "".join(del_tokens[:prefix_len])))

            end_del = len(del_tokens) - suffix_len if suffix_len else len(del_tokens)
            end_ins = len(ins_tokens) - suffix_len if suffix_len else len(ins_tokens)
            mid_del = del_tokens[prefix_len:end_del]
            mid_ins = ins_tokens[prefix_len:end_ins]

            if mid_del:
                refined.append((DELETE, "".join(mid_del)))
            if mid_ins:
                refined.append((INSERT, "".join(mid_ins)))

            if suffix_len > 0:
                refined.append((EQUAL, "".join(del_tokens[end_del:])))

            i += 2
        else:
            refined.append(diffs[i])
            i += 1
    return refined


def _merge_adjacent_diffs(diffs: List[tuple]) -> List[tuple]:
    """Merge consecutive diff entries that share the same operation."""
    if not diffs:
        return diffs
    merged = [diffs[0]]
    for op, text in diffs[1:]:
        if op == merged[-1][0]:
            merged[-1] = (op, merged[-1][1] + text)
        else:
            merged.append((op, text))
    return merged


def _iso_now_utc() -> str:
    """Return current UTC time in Word-friendly ISO8601 format."""
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z")


def _ensure_review_settings(doc: Document) -> None:
    """
    Force review metadata/settings so Word shows tracked revisions and
    comment balloons by default when the document is opened.
    """
    settings = doc.settings.element

    if settings.find(qn("w:trackRevisions")) is None:
        settings.append(OxmlElement("w:trackRevisions"))

    revision_view = settings.find(qn("w:revisionView"))
    if revision_view is None:
        revision_view = OxmlElement("w:revisionView")
        settings.append(revision_view)

    revision_view.set(qn("w:markup"), "1")
    revision_view.set(qn("w:comments"), "1")
    revision_view.set(qn("w:insDel"), "1")
    revision_view.set(qn("w:formatting"), "1")
    revision_view.set(qn("w:inkAnnotations"), "1")

    # Show comments and formatting changes in balloons (margin bubbles).
    # <w:showComments/> and <w:showInsDel/> are Word 2010+ elements that
    # control whether balloons appear in the margin for comments/revisions.
    for tag in ("w:showComments", "w:showInsDel"):
        existing = settings.find(qn(tag))
        if existing is None:
            elm = OxmlElement(tag)
            elm.set(qn("w:val"), "1")
            settings.append(elm)

    for tag in (
        "w:doNotShowInsDel",
        "w:doNotShowMarkup",
        "w:doNotShowComments",
    ):
        node = settings.find(qn(tag))
        if node is not None:
            settings.remove(node)


def _needs_space_preserve(text: str) -> bool:
    return bool(text) and (text[0].isspace() or text[-1].isspace())


def _append_tracked_insert(
    para: Paragraph,
    text: str,
    revision_id: int,
    author: str,
    date_iso: str,
) -> Optional[Run]:
    """Append a native Word insertion revision and return run anchor."""
    if not text:
        return None
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), str(revision_id))
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), date_iso)

    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    if _needs_space_preserve(text):
        t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    ins.append(r)
    para._p.append(ins)
    return Run(r, para)


def _append_tracked_delete(
    para: Paragraph,
    text: str,
    revision_id: int,
    author: str,
    date_iso: str,
) -> Optional[Run]:
    """Append a native Word deletion revision and return run anchor."""
    if not text:
        return None
    deleted = OxmlElement("w:del")
    deleted.set(qn("w:id"), str(revision_id))
    deleted.set(qn("w:author"), author)
    deleted.set(qn("w:date"), date_iso)

    r = OxmlElement("w:r")
    del_text = OxmlElement("w:delText")
    if _needs_space_preserve(text):
        del_text.set(qn("xml:space"), "preserve")
    del_text.text = text
    r.append(del_text)
    deleted.append(r)
    para._p.append(deleted)
    return Run(r, para)


def _find_mod_reason(
    original_para: str,
    modified_para: str,
    changed_text: str,
    modifications: List[Dict[str, Any]],
    revisions: Optional[List[Dict[str, Any]]] = None,
) -> str:
    """Pick the most relevant playbook rationale for a changed paragraph.

    Prioritizes revision-level reasoning (from v2 pipeline) which contains
    rich AI-generated explanations, falling back to modification-level
    explanations (legacy).
    """
    if revisions:
        reason = _find_revision_reason(original_para, modified_para, revisions)
        if reason:
            return reason

    changed_norm = changed_text.strip()
    best_reason = ""

    for mod in modifications:
        orig_fragment = str(
            mod.get("original_fragment")
            or mod.get("original_text")
            or mod.get("source_fragment")
            or mod.get("from")
            or ""
        )
        mod_fragment = str(
            mod.get("modified_fragment")
            or mod.get("modified_text")
            or mod.get("replacement")
            or mod.get("replace_with")
            or mod.get("to")
            or ""
        )
        explanation = str(
            mod.get("explanation")
            or mod.get("reason")
            or mod.get("rationale")
            or mod.get("ai_reason")
            or mod.get("modification_reason")
            or ""
        ).strip()
        rule_title = str(mod.get("rule_title", "") or "").strip()
        rule_id = str(mod.get("rule_id", "") or "").strip()
        mod_type = str(mod.get("modification_type", "") or "").strip()

        parts: List[str] = []
        if rule_id:
            parts.append(f"[{rule_id}]")
        if rule_title:
            parts.append(rule_title)

        if explanation:
            reason_lines = [
                "AI审阅修改说明",
                f"规则依据: {' '.join(parts).strip()}".strip(),
                f"AI判断: {explanation}",
            ]
        else:
            reason_lines = [
                "AI审阅修改说明",
                f"规则依据: {' '.join(parts).strip() or '已选 Playbook 规则'}",
                "AI判断: 该处表述与已选规则不一致，需要进行定向修订。",
            ]

        if orig_fragment or mod_fragment:
            reason_lines.append(
                f"修改内容: \"{_short_reason(orig_fragment, max_len=160) or '(原文片段)'}\" -> "
                f"\"{_short_reason(mod_fragment, max_len=160) or '(新文本片段)'}\""
            )

        if mod_type:
            reason_lines.append(f"修改类型: {mod_type}")

        reason = _short_reason("\n".join(reason_lines), max_len=1200)

        if changed_norm and (
            (orig_fragment and changed_norm in orig_fragment)
            or (mod_fragment and changed_norm in mod_fragment)
        ):
            return reason

        if (
            (orig_fragment and orig_fragment in original_para)
            or (mod_fragment and mod_fragment in modified_para)
        ):
            if not best_reason:
                best_reason = reason

    return best_reason


def _find_revision_reason(
    original_para: str,
    modified_para: str,
    revisions: List[Dict[str, Any]],
) -> str:
    """Find the reasoning from a revision that matches this paragraph.

    Produces a structured comment with rule IDs, reasoning, and individual
    change descriptions so reviewers can understand *why* each edit was made.
    """
    norm_para = _normalize_paragraph_for_alignment(original_para)

    for rev in revisions:
        orig_clause = rev.get("original_clause", "")
        if not orig_clause:
            continue

        norm_clause = _normalize_paragraph_for_alignment(orig_clause)
        if not norm_clause or not norm_para:
            continue

        if norm_clause in norm_para or norm_para in norm_clause:
            return _build_revision_comment(rev)

    return ""


def _build_revision_comment(rev: Dict[str, Any]) -> str:
    """Build a structured comment string from a revision dict."""
    parts: List[str] = ["AI审阅修改说明"]

    rule_ids = rev.get("applicable_rule_ids", [])
    if rule_ids:
        label = ", ".join(rule_ids) if isinstance(rule_ids, list) else str(rule_ids)
        parts.append(f"规则依据: {label}")

    reasoning = (rev.get("reasoning") or "").strip()
    if reasoning:
        parts.append(f"AI判断: {reasoning}")

    changes = rev.get("changes_made", [])
    if changes:
        parts.append("修改明细:")
        for ch in changes:
            what = (ch.get("what") or "").strip()
            why = (ch.get("why") or "").strip()
            rule_id = (ch.get("rule_id") or "").strip()
            if what and why:
                line = f"• 修改: {what}；原因: {why}"
                if rule_id:
                    line += f"（{rule_id}）"
                parts.append(line)
            elif what:
                parts.append(f"• 修改: {what}")

    if len(parts) == 1:
        parts.append("规则依据: 已选 Playbook 规则")
        parts.append("AI判断: 该处条款与规则要求不一致，已按规则进行最小必要修改。")

    return _short_reason("\n".join(parts), max_len=1200)


def _try_add_comment(doc: Document, anchor_run: Run, reason: str) -> None:
    """Best-effort add_comment with multi-paragraph support.

    If *reason* contains newlines, each line becomes a separate paragraph
    inside the comment so that rule IDs, reasoning, and change bullets are
    visually distinct in the Word comment pane.
    """
    if not reason:
        return
    try:
        # Write full rationale in one call. This is more reliable across
        # python-docx versions than trying to append paragraphs later.
        doc.add_comment(
            anchor_run,
            text=reason,
            author="AI Legal Assistant",
            initials="AI",
        )
    except Exception:
        try:
            doc.add_comment(
                anchor_run,
                text=_short_reason(reason, max_len=300),
                author="AI Legal Assistant",
                initials="AI",
            )
        except Exception:
            pass


def _append_comment_paragraph(comment_element, text: str) -> None:
    """Append an additional <w:p> to an existing comment element."""
    try:
        p = OxmlElement("w:p")
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        if _needs_space_preserve(text):
            t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        p.append(r)
        comment_element.append(p)
    except Exception:
        pass


def _add_redline_paragraphs(
    doc: Document,
    original: str,
    modified: str,
    modifications: List[Dict[str, Any]],
    revisions: Optional[List[Dict[str, Any]]] = None,
):
    """
    Split both texts into paragraphs, diff each pair, and write
    redline-formatted runs into the Word document.
    """
    orig_paras = original.split("\n")
    mod_paras = modified.split("\n")

    max_len = max(len(orig_paras), len(mod_paras))

    while len(orig_paras) < max_len:
        orig_paras.append("")
    while len(mod_paras) < max_len:
        mod_paras.append("")

    author = "AI Legal Assistant"
    revision_id = 1
    date_iso = _iso_now_utc()

    for orig_p, mod_p in zip(orig_paras, mod_paras):
        para = doc.add_paragraph()
        revision_id = _render_paragraph_diff(
            doc=doc,
            para=para,
            orig_p=orig_p,
            mod_p=mod_p,
            modifications=modifications,
            author=author,
            revision_id=revision_id,
            date_iso=date_iso,
            revisions=revisions,
        )


def _clear_paragraph_content(para: Paragraph) -> None:
    """Remove runs/revisions while preserving paragraph properties (w:pPr)."""
    p_elm = para._p
    for child in list(p_elm):
        if child.tag != qn("w:pPr"):
            p_elm.remove(child)


def _normalize_paragraph_for_alignment(text: str) -> str:
    if not text:
        return ""
    return re.sub(r"\s+", " ", text.translate(_PARA_PUNCT_MAP)).strip()


def _render_paragraph_diff(
    doc: Document,
    para: Paragraph,
    orig_p: str,
    mod_p: str,
    modifications: List[Dict[str, Any]],
    author: str,
    revision_id: int,
    date_iso: str,
    revisions: Optional[List[Dict[str, Any]]] = None,
) -> int:
    """
    Render a paragraph diff into an existing paragraph and return next revision id.

    - If normalized text is effectively equal, keep paragraph clean.
    - Only changed paragraphs get rewritten with ins/del revisions.
    """
    _clear_paragraph_content(para)

    if orig_p == mod_p or (
        _normalize_paragraph_for_alignment(orig_p)
        == _normalize_paragraph_for_alignment(mod_p)
    ):
        if mod_p:
            run = para.add_run(mod_p)
            run.font.color.rgb = _BLACK
        return revision_id

    diffs = compute_word_diffs(orig_p, mod_p)
    comment_anchor: Optional[Run] = None
    comment_reason = ""
    for op, text in diffs:
        if not text:
            continue
        if op == dmp_module.diff_match_patch.DIFF_EQUAL:
            run = para.add_run(text)
            run.font.color.rgb = _BLACK
            if comment_anchor is None:
                comment_anchor = run
        elif op == dmp_module.diff_match_patch.DIFF_DELETE:
            deleted_run = _append_tracked_delete(
                para=para,
                text=text,
                revision_id=revision_id,
                author=author,
                date_iso=date_iso,
            )
            revision_id += 1
            if comment_anchor is None and deleted_run is not None:
                comment_anchor = deleted_run
            if not comment_reason:
                comment_reason = _find_mod_reason(orig_p, mod_p, text, modifications, revisions=revisions)
        elif op == dmp_module.diff_match_patch.DIFF_INSERT:
            inserted_run = _append_tracked_insert(
                para=para,
                text=text,
                revision_id=revision_id,
                author=author,
                date_iso=date_iso,
            )
            revision_id += 1
            if comment_anchor is None and inserted_run is not None:
                comment_anchor = inserted_run
            if not comment_reason:
                comment_reason = _find_mod_reason(orig_p, mod_p, text, modifications, revisions=revisions)

    if not comment_reason:
        comment_reason = "Playbook-driven edit: aligned clause with selected review rules."
    if comment_anchor is not None:
        _try_add_comment(doc, comment_anchor, comment_reason)
    return revision_id


def _split_contract_paragraphs(text: str) -> List[str]:
    normalized = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    if not normalized.strip():
        return []
    if "\n\n" in normalized:
        chunks = re.split(r"\n{2,}", normalized)
    else:
        chunks = normalized.split("\n")
    return [p.strip() for p in chunks if p.strip()]


def _get_body_blocks_with_paragraphs(doc: Document) -> List[tuple]:
    """
    Return [(Paragraph, block_text), ...] in document order, including content
    from sdt (content controls). Aligns with extraction that includes sdt.
    """
    body = doc.element.body
    blocks: List[tuple] = []
    para_to_elem = {id(p._p): p for p in doc.paragraphs}

    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            texts = []
            for t in child.iter():
                if t.tag.endswith("}t") and t.text:
                    texts.append(t.text)
            block_text = "".join(texts).strip()
            if block_text:
                para = para_to_elem.get(id(child))
                if para is not None:
                    blocks.append((para, block_text))
        elif tag == "sdt":
            sdt_content = None
            for elem in child.iter():
                if elem.tag.split("}")[-1] == "sdtContent":
                    sdt_content = elem
                    break
            if sdt_content is not None:
                for p_elem in sdt_content.iter():
                    if p_elem.tag.endswith("}p"):
                        texts = []
                        for t in p_elem.iter():
                            if t.tag.endswith("}t") and t.text:
                                texts.append(t.text)
                        block_text = "".join(texts).strip()
                        if block_text:
                            parent = p_elem.getparent()
                            if parent is not None:
                                para = Paragraph(p_elem, parent)
                                blocks.append((para, block_text))
                        break

    return blocks


def _add_redline_paragraphs_from_template(
    doc: Document,
    original: str,
    modified: str,
    modifications: List[Dict[str, Any]],
    revisions: Optional[List[Dict[str, Any]]] = None,
):
    """
    Rewrite the source document's existing paragraphs so paragraph-level formatting
    (including native numbering) is preserved in output.
    """
    orig_paras = _split_contract_paragraphs(original)
    mod_paras = _split_contract_paragraphs(modified)
    if not orig_paras and not mod_paras:
        return

    # Build existing_paras in same order as extraction (including sdt content controls)
    blocks = _get_body_blocks_with_paragraphs(doc)
    existing_paras = [p for p, _ in blocks]
    while len(existing_paras) < len(orig_paras):
        existing_paras.append(doc.add_paragraph())

    author = "AI Legal Assistant"
    revision_id = 1
    date_iso = _iso_now_utc()
    matcher = SequenceMatcher(
        a=[_normalize_paragraph_for_alignment(p) for p in orig_paras],
        b=[_normalize_paragraph_for_alignment(p) for p in mod_paras],
        autojunk=False,
    )

    def _insert_before(anchor: Optional[Paragraph], style_source: Optional[Paragraph]) -> Paragraph:
        if anchor is not None:
            p = anchor.insert_paragraph_before("")
        else:
            p = doc.add_paragraph()
        if style_source is not None:
            try:
                p.style = style_source.style
            except Exception:
                pass
        return p

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            # Keep untouched paragraphs as-is to preserve run-level formatting.
            continue

        if tag == "replace":
            pair_count = min(i2 - i1, j2 - j1)
            for offset in range(pair_count):
                orig_idx = i1 + offset
                mod_idx = j1 + offset
                para = existing_paras[orig_idx]
                revision_id = _render_paragraph_diff(
                    doc=doc,
                    para=para,
                    orig_p=orig_paras[orig_idx],
                    mod_p=mod_paras[mod_idx],
                    modifications=modifications,
                    author=author,
                    revision_id=revision_id,
                    date_iso=date_iso,
                    revisions=revisions,
                )

            for orig_idx in range(i1 + pair_count, i2):
                para = existing_paras[orig_idx]
                revision_id = _render_paragraph_diff(
                    doc=doc,
                    para=para,
                    orig_p=orig_paras[orig_idx],
                    mod_p="",
                    modifications=modifications,
                    author=author,
                    revision_id=revision_id,
                    date_iso=date_iso,
                    revisions=revisions,
                )

            if j1 + pair_count < j2:
                anchor_idx = i1 + pair_count
                anchor = existing_paras[anchor_idx] if anchor_idx < len(existing_paras) else None
                style_source = (
                    anchor
                    if anchor is not None
                    else (existing_paras[anchor_idx - 1] if anchor_idx > 0 else None)
                )
                for mod_idx in range(j1 + pair_count, j2):
                    para = _insert_before(anchor=anchor, style_source=style_source)
                    revision_id = _render_paragraph_diff(
                        doc=doc,
                        para=para,
                        orig_p="",
                        mod_p=mod_paras[mod_idx],
                        modifications=modifications,
                        author=author,
                        revision_id=revision_id,
                        date_iso=date_iso,
                        revisions=revisions,
                    )
            continue

        if tag == "delete":
            for orig_idx in range(i1, i2):
                para = existing_paras[orig_idx]
                revision_id = _render_paragraph_diff(
                    doc=doc,
                    para=para,
                    orig_p=orig_paras[orig_idx],
                    mod_p="",
                    modifications=modifications,
                    author=author,
                    revision_id=revision_id,
                    date_iso=date_iso,
                    revisions=revisions,
                )
            continue

        if tag == "insert":
            anchor = existing_paras[i1] if i1 < len(existing_paras) else None
            style_source = anchor if anchor is not None else (existing_paras[i1 - 1] if i1 > 0 else None)
            for mod_idx in range(j1, j2):
                para = _insert_before(anchor=anchor, style_source=style_source)
                revision_id = _render_paragraph_diff(
                    doc=doc,
                    para=para,
                    orig_p="",
                    mod_p=mod_paras[mod_idx],
                    modifications=modifications,
                    author=author,
                    revision_id=revision_id,
                    date_iso=date_iso,
                    revisions=revisions,
                )


_SIGNATURE_MARKERS = (
    "Very truly yours",
    "Sincerely",
    "IN WITNESS WHEREOF",
    "(Remainder of page intentionally left blank)",
)


def _sort_insertions_for_mandatory_order(insertions: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Sort insertions so mandatory_language clauses appear in playbook order:
    A (Dual Representative) first, B (Non-Restriction) second, C (Data Room) third.
    Blind NDA (if present) comes before the three mandatory clauses.
    """
    def _order_key(ins: Dict[str, Any]) -> int:
        text = (ins.get("clause_text") or "").lower()
        if "identity disclosure" in text or "conflict check" in text or "conflicts check" in text:
            return 0  # Blind NDA first
        if "dual representative" in text:
            return 1  # Rule A
        if "investment business" in text and "competitors" in text:
            return 2  # Rule B
        if "electronic data room" in text and ("modified" in text or "amended" in text):
            return 3  # Rule C
        return 4  # Other insertions last

    return sorted(insertions, key=_order_key)


def _find_insertion_anchor(doc: Document) -> Optional[Paragraph]:
    """
    Find the anchor for insertions: the paragraph immediately AFTER the last
    numbered body paragraph. Inserting before this keeps numbering continuous.
    """
    paras = list(doc.paragraphs)
    last_num_idx = -1
    for i, p in enumerate(paras):
        pPr = p._p.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:numPr")) is not None:
            last_num_idx = i

    if last_num_idx < 0:
        for para in paras:
            if "Very truly yours" in (para.text or ""):
                return para
        return None

    next_idx = last_num_idx + 1
    if next_idx < len(paras):
        return paras[next_idx]
    return paras[last_num_idx] if last_num_idx >= 0 else None


def _find_signature_anchor(doc: Document) -> Optional[Paragraph]:
    """Find the first paragraph that contains a signature/closing marker."""
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        for marker in _SIGNATURE_MARKERS:
            if marker in text:
                return para
    return None


def _get_style_and_num_pr(doc: Document, anchor: Optional[Paragraph]) -> tuple:
    """Get style and numPr from the last numbered body paragraph before the anchor."""
    paras = list(doc.paragraphs)
    start_idx = len(paras) - 1
    if anchor is not None:
        for idx, p in enumerate(paras):
            if p is anchor:
                start_idx = idx - 1
                break
    for i in range(start_idx, -1, -1):
        p = paras[i]
        p_elm = p._p
        pPr = p_elm.find(qn("w:pPr"))
        if pPr is None:
            continue
        numPr = pPr.find(qn("w:numPr"))
        if numPr is not None:
            return p.style, numPr
    return None, None


def _apply_paragraph_format(para: Paragraph, style_source, num_pr_source) -> None:
    """Apply style and numbering to a paragraph to match the document body."""
    p_elm = para._p
    pPr = p_elm.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_elm.insert(0, pPr)
    if style_source is not None:
        try:
            para.style = style_source
        except Exception:
            pass
    if num_pr_source is not None:
        existing_num = pPr.find(qn("w:numPr"))
        if existing_num is not None:
            pPr.remove(existing_num)
        pPr.append(num_pr_source)


def _add_insertion_paragraphs(doc: Document, insertions: List[Dict[str, Any]]) -> None:
    """
    Append Step 2b insertion clauses as pure tracked-insert paragraphs,
    placed immediately after the last numbered body paragraph for continuous
    numbering. Order: Blind NDA (if any), then mandatory_language A, B, C.
    """
    if not insertions:
        return

    insertions = _sort_insertions_for_mandatory_order(insertions)
    anchor = _find_insertion_anchor(doc)
    if anchor is None:
        anchor = _find_signature_anchor(doc)
    style_source, num_pr_source = _get_style_and_num_pr(doc, anchor)
    author = "AI Legal Assistant"
    date_iso = _iso_now_utc()
    revision_id = 9000

    for ins in reversed(insertions):
        clause_text = (ins.get("clause_text") or "").strip()
        if not clause_text:
            continue

        chunks = [c.strip() for c in clause_text.split("\n\n") if c.strip()]

        # When inserting before an anchor, each insert_paragraph_before +
        # anchor update reverses iteration order (same reason the outer loop
        # uses reversed()).  Reverse chunks so they end up in source order.
        # When appending (anchor is None), forward order is already correct.
        ordered_chunks = list(reversed(chunks)) if anchor is not None else chunks

        for chunk in ordered_chunks:
            if anchor is not None:
                para = anchor.insert_paragraph_before("")
            else:
                para = doc.add_paragraph()

            _clear_paragraph_content(para)
            if num_pr_source is not None:
                _apply_paragraph_format(para, style_source, copy.deepcopy(num_pr_source))
            elif style_source is not None:
                try:
                    para.style = style_source
                except Exception:
                    pass

            _append_tracked_insert(
                para=para,
                text=chunk,
                revision_id=revision_id,
                author=author,
                date_iso=date_iso,
            )
            revision_id += 1

            if anchor is not None:
                anchor = para


def _short_reason(reason: str, max_len: int = 500) -> str:
    """Keep rationale readable for Word comments while preserving enough detail.

    Preserves intentional newlines (used for multi-paragraph comments) while
    collapsing redundant whitespace within each line.
    """
    if not reason:
        return ""
    lines = reason.split("\n")
    compact_lines = [" ".join(line.split()) for line in lines]
    compact = "\n".join(line for line in compact_lines if line)
    if len(compact) <= max_len:
        return compact
    return compact[: max_len - 3].rstrip() + "..."


def _add_issues_table(doc: Document, issues: List[Dict[str, Any]]):
    """Render the issues list as a formatted table."""
    headers = ["#", "Severity", "Category", "Issue", "Clause", "Status"]
    table = doc.add_table(rows=1, cols=len(headers), style="Light Grid Accent 1")

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True

    # Data rows
    for issue in issues:
        row = table.add_row()
        row.cells[0].text = str(issue.get("id", ""))
        row.cells[1].text = str(issue.get("severity", ""))
        row.cells[2].text = str(issue.get("category", ""))
        row.cells[3].text = str(issue.get("title", issue.get("description", "")))
        row.cells[4].text = str(issue.get("clause_reference", ""))
        row.cells[5].text = str(issue.get("status", ""))

    # Severity colouring
    severity_colors = {"P0": _RED, "P1": RGBColor(0xE6, 0x7E, 0x22), "P2": _GREY}
    for row_idx in range(1, len(table.rows)):
        sev_cell = table.rows[row_idx].cells[1]
        sev_text = sev_cell.text.strip()
        if sev_text in severity_colors:
            for para in sev_cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = severity_colors[sev_text]
                    run.bold = True


def generate_clean_docx(
    text: str,
    title: str = "Contract — Clean Copy",
) -> io.BytesIO:
    """Generate a clean Word document (no redlines) with the modified text."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading(title, level=1)
    for para_text in text.split("\n\n"):
        stripped = para_text.strip()
        if stripped:
            doc.add_paragraph(stripped)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
