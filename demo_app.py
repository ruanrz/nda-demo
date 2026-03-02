"""
AI Legal Assistant — Contract Review Demo

Architecture (v3):
  - Unified pipeline: analysis (LLM) → execution (local + LLM fallback) → issues (LLM)
  - Model preset: quality (o3 + gpt-4o)
  - Word Track Changes output
  - Issues List / Risk Summary
  - Own Paper + Counterparty Paper modes
  - Rule Learning from historical contracts (OpenAI provider)
"""

import os
import sys
import json
import logging
import re
from datetime import datetime
from io import BytesIO
import streamlit as st
import diff_match_patch as dmp_module
from docx import Document
from docx.oxml.ns import qn

# ── Logging ──────────────────────────────────────────────────────
LOG_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "logs")
os.makedirs(LOG_DIR, exist_ok=True)
LOG_FORMAT = "%(asctime)s | %(levelname)-8s | %(name)s | %(funcName)s:%(lineno)d | %(message)s"
LOG_DATE_FORMAT = "%Y-%m-%d %H:%M:%S"

logger = logging.getLogger("legal_ai_demo")
logger.setLevel(logging.DEBUG)
if not logger.handlers:
    ch = logging.StreamHandler()
    ch.setLevel(logging.INFO)
    ch.setFormatter(logging.Formatter(LOG_FORMAT, LOG_DATE_FORMAT))
    logger.addHandler(ch)
    fh = logging.FileHandler(
        os.path.join(LOG_DIR, f"demo_{datetime.now().strftime('%Y%m%d')}.log"),
        encoding="utf-8",
    )
    fh.setLevel(logging.DEBUG)
    fh.setFormatter(logging.Formatter(LOG_FORMAT, LOG_DATE_FORMAT))
    logger.addHandler(fh)

# ── Path setup ───────────────────────────────────────────────────
_current_dir = os.path.dirname(os.path.abspath(__file__))
_providers_dir = os.path.join(_current_dir, "providers")
for p in (_providers_dir, _current_dir):
    if p not in sys.path:
        sys.path.insert(0, p)

# ── Import: NEW unified review pipeline ─────────────────────────
try:
    from providers.unified_review import unified_review_contract
    from providers.word_generator import generate_redline_docx, generate_clean_docx
    from providers.playbook_loader import (
        load_playbooks_from_markdown,
        load_playbooks_for_display,
    )
    UNIFIED_AVAILABLE = True
    logger.info("Unified review pipeline loaded (OpenAI)")
except Exception as e:
    UNIFIED_AVAILABLE = False
    logger.warning(f"Unified pipeline unavailable: {e}")

# ── Import: Rule-learning provider (OpenAI) ───────────────────────
try:
    from providers.rule_learning import (
        learn_from_contract_diff,
        load_learned_rules,
        save_learned_rules,
        get_learned_rules_stats,
        LearnedRule,
        LearnedRulesStore,
    )
    LEARNING_AVAILABLE = True
    logger.info("Rule learning provider loaded (OpenAI)")
except ImportError:
    LEARNING_AVAILABLE = False
    logger.warning("Rule learning provider unavailable")

# ── Constants ────────────────────────────────────────────────────
PLAYBOOK_FINAL_FILE = os.path.join(_current_dir, "playbook_mapping_final.json")

# ═══════════════════════════════════════════════════════════════════
# Streamlit App
# ═══════════════════════════════════════════════════════════════════

st.set_page_config(layout="wide", page_title="AI Legal Assistant", initial_sidebar_state="expanded")

st.markdown("""
<style>
    .redline-add {
        background-color: #e6ffec; color: #1a7f37;
        font-weight: bold; padding: 0 2px; border-radius: 2px;
    }
    .redline-del {
        background-color: #ffebe9; color: #cf222e;
        text-decoration: line-through; opacity: 0.85;
        padding: 0 2px; border-radius: 2px;
    }
    .clause-box {
        border: 1px solid #e1e4e8; padding: 15px;
        border-radius: 6px; margin-bottom: 10px; background: white;
    }
    .issue-p0 { border-left: 4px solid #cf222e; padding-left: 12px; }
    .issue-p1 { border-left: 4px solid #e67e22; padding-left: 12px; }
    .issue-p2 { border-left: 4px solid #999;    padding-left: 12px; }
    .stat-card {
        background: #f8f9fa; border-radius: 8px;
        padding: 16px; text-align: center;
    }
    .stCodeBlock code { white-space: pre-wrap !important; word-wrap: break-word !important; }
    .stCodeBlock pre  { white-space: pre-wrap !important; }
    [data-testid="stSidebar"] { min-width: 460px; width: 460px; }
    [data-testid="stSidebar"] .block-container {
        padding-top: 1rem;
    }
    [data-testid="stSidebar"] .stExpander {
        border-radius: 8px;
        border: 1px solid #e1e4e8;
    }
</style>
""", unsafe_allow_html=True)


# ── Helper functions ─────────────────────────────────────────────

def generate_diff_html(text1: str, text2: str) -> str:
    if not text1 or not text2:
        return ""
    dmp = dmp_module.diff_match_patch()
    dmp.Diff_Timeout = 1.5
    diffs = dmp.diff_main(text1, text2)
    dmp.diff_cleanupSemantic(diffs)
    html = ""
    for op, data in diffs:
        escaped = data.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if op == dmp_module.diff_match_patch.DIFF_INSERT:
            html += f"<span class='redline-add'>{escaped}</span>"
        elif op == dmp_module.diff_match_patch.DIFF_DELETE:
            html += f"<span class='redline-del'>{escaped}</span>"
        else:
            html += f"<span>{escaped}</span>"
    return html


def diff_changed(text1: str, text2: str):
    if text1 == text2:
        return False, 0
    dmp = dmp_module.diff_match_patch()
    dmp.Diff_Timeout = 0.5
    diffs = dmp.diff_main(text1, text2)
    dmp.diff_cleanupSemantic(diffs)
    ops = [d for d in diffs if d[0] != dmp_module.diff_match_patch.DIFF_EQUAL]
    return len(ops) > 0, len(ops)


_NUM_TOKEN_RE = re.compile(r"%(\d+)")


def _to_roman(value: int) -> str:
    if value <= 0:
        return str(value)
    pairs = [
        (1000, "M"),
        (900, "CM"),
        (500, "D"),
        (400, "CD"),
        (100, "C"),
        (90, "XC"),
        (50, "L"),
        (40, "XL"),
        (10, "X"),
        (9, "IX"),
        (5, "V"),
        (4, "IV"),
        (1, "I"),
    ]
    result = []
    remaining = value
    for number, symbol in pairs:
        while remaining >= number:
            result.append(symbol)
            remaining -= number
    return "".join(result)


def _to_alpha(value: int, uppercase: bool = False) -> str:
    if value <= 0:
        return str(value)
    letters = []
    n = value
    while n > 0:
        n -= 1
        letters.append(chr(ord("A" if uppercase else "a") + (n % 26)))
        n //= 26
    return "".join(reversed(letters))


def _format_number_token(value: int, num_fmt: str) -> str:
    fmt = (num_fmt or "decimal").lower()
    if fmt == "decimalzero" and value < 10:
        return f"0{value}"
    if fmt in {"decimal", "decimalzero"}:
        return str(value)
    if fmt == "lowerletter":
        return _to_alpha(value, uppercase=False)
    if fmt == "upperletter":
        return _to_alpha(value, uppercase=True)
    if fmt == "lowerroman":
        return _to_roman(value).lower()
    if fmt == "upperroman":
        return _to_roman(value)
    return str(value)


def _parse_level_definition(level_elm):
    num_fmt_elm = level_elm.find(qn("w:numFmt"))
    lvl_text_elm = level_elm.find(qn("w:lvlText"))
    start_elm = level_elm.find(qn("w:start"))
    num_fmt = num_fmt_elm.get(qn("w:val")) if num_fmt_elm is not None else "decimal"
    lvl_text = lvl_text_elm.get(qn("w:val")) if lvl_text_elm is not None else None
    start_val = 1
    if start_elm is not None:
        raw_start = start_elm.get(qn("w:val"))
        if raw_start is not None:
            try:
                start_val = int(raw_start)
            except ValueError:
                start_val = 1
    return {
        "num_fmt": num_fmt or "decimal",
        "lvl_text": lvl_text,
        "start": start_val,
    }


def _load_numbering_levels(doc: Document):
    try:
        numbering_root = doc.part.numbering_part.element
    except Exception:
        return {}

    abstract_levels = {}
    for abstract in numbering_root.findall(qn("w:abstractNum")):
        abstract_id = abstract.get(qn("w:abstractNumId"))
        if not abstract_id:
            continue
        level_map = {}
        for level in abstract.findall(qn("w:lvl")):
            ilvl_raw = level.get(qn("w:ilvl"), "0")
            try:
                ilvl = int(ilvl_raw)
            except ValueError:
                ilvl = 0
            level_map[ilvl] = _parse_level_definition(level)
        abstract_levels[abstract_id] = level_map

    numbering_levels = {}
    for num in numbering_root.findall(qn("w:num")):
        num_id = num.get(qn("w:numId"))
        if not num_id:
            continue
        abs_elm = num.find(qn("w:abstractNumId"))
        if abs_elm is None:
            continue
        abstract_id = abs_elm.get(qn("w:val"))
        if not abstract_id:
            continue

        merged_levels = dict(abstract_levels.get(abstract_id, {}))
        for override in num.findall(qn("w:lvlOverride")):
            ilvl_raw = override.get(qn("w:ilvl"), "0")
            try:
                ilvl = int(ilvl_raw)
            except ValueError:
                ilvl = 0

            override_level = override.find(qn("w:lvl"))
            if override_level is not None:
                merged_levels[ilvl] = _parse_level_definition(override_level)

            start_override = override.find(qn("w:startOverride"))
            if start_override is not None:
                raw_start = start_override.get(qn("w:val"))
                if raw_start is not None:
                    try:
                        start_val = int(raw_start)
                    except ValueError:
                        start_val = 1
                    level_def = dict(merged_levels.get(ilvl, {}))
                    level_def["start"] = start_val
                    level_def.setdefault("num_fmt", "decimal")
                    merged_levels[ilvl] = level_def

        numbering_levels[num_id] = merged_levels

    return numbering_levels


def _extract_docx_text_with_numbering(content: bytes) -> str:
    doc = Document(BytesIO(content))
    numbering_levels = _load_numbering_levels(doc)
    list_counters = {}
    extracted = []

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue

        prefix = ""
        ppr = para._p.pPr
        num_pr = ppr.numPr if ppr is not None else None
        if num_pr is not None and num_pr.numId is not None:
            num_id = str(num_pr.numId.val)
            try:
                ilvl = int(num_pr.ilvl.val) if num_pr.ilvl is not None else 0
            except ValueError:
                ilvl = 0

            level_defs = numbering_levels.get(num_id, {})
            level_state = list_counters.setdefault(num_id, {})

            start_val = int(level_defs.get(ilvl, {}).get("start", 1))
            current = level_state.get(ilvl, start_val - 1) + 1
            level_state[ilvl] = current

            for key in list(level_state.keys()):
                if key > ilvl:
                    level_state.pop(key, None)

            level_def = level_defs.get(ilvl, {})
            num_fmt = str(level_def.get("num_fmt", "decimal"))
            lvl_text = level_def.get("lvl_text") or f"%{ilvl + 1}."

            if num_fmt.lower() == "bullet":
                prefix = "•"
            else:
                def _replace_level_token(match):
                    level_index = int(match.group(1)) - 1
                    level_value = level_state.get(level_index)
                    if level_value is None:
                        return ""
                    token_fmt = str(level_defs.get(level_index, {}).get("num_fmt", "decimal"))
                    return _format_number_token(level_value, token_fmt)

                prefix = _NUM_TOKEN_RE.sub(_replace_level_token, str(lvl_text))
                prefix = prefix.replace("\t", " ").strip()
                if not prefix:
                    prefix = f"{current}."

        if prefix and not text.startswith(prefix):
            text = f"{prefix} {text}"
        extracted.append(text)

    return "\n\n".join(extracted)


def parse_uploaded_contract(uploaded_file, store_source_docx: bool = False):
    if not uploaded_file:
        if store_source_docx:
            st.session_state.pop("review_source_docx_bytes", None)
        return None
    name = uploaded_file.name.lower()
    if name.endswith(".docx"):
        try:
            content = uploaded_file.read()
            if store_source_docx:
                st.session_state["review_source_docx_bytes"] = content
            doc = Document(BytesIO(content))
            return "\n\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            if store_source_docx:
                st.session_state.pop("review_source_docx_bytes", None)
            st.warning(f".docx parsing failed: {e}")
            return None
    elif name.endswith(".txt"):
        if store_source_docx:
            st.session_state.pop("review_source_docx_bytes", None)
        return uploaded_file.read().decode("utf-8")
    if store_source_docx:
        st.session_state.pop("review_source_docx_bytes", None)
    st.warning("Only .docx and .txt files are supported.")
    return None


def _friendly_review_error_message(err: Exception) -> str:
    """Map common provider failures to actionable user-facing guidance."""
    text = str(err)
    lower = text.lower()
    if "unsupported_country_region_territory" in lower:
        return (
            "OpenAI 请求被地区策略拒绝。请配置可访问的 OpenAI 兼容 endpoint（设置 "
            "`OPENAI_API_BASE` 或 `OPENAI_BASE_URL`），并确保所选模型在该 endpoint 可用。"
        )
    if "invalid_api_key" in lower or "authentication" in lower:
        return "OpenAI 鉴权失败。请检查 `OPENAI_API_KEY` 是否正确并重新启动应用。"
    if "insufficient_quota" in lower or "quota" in lower:
        return "OpenAI 配额不足或账单受限。请检查账号额度与 billing 状态。"
    return text


def build_default_contract() -> str:
    return """CONFIDENTIALITY AGREEMENT

1. DEFINITIONS
"Representatives" means the Recipient's directors, officers and employees.

2. CONFIDENTIAL INFORMATION
"Confidential Information" means any and all technical, non-public, and non-technical information provided by the Company to the Recipient, which may include without limitation information regarding (a) patent and patent applications; (b) trade secrets; and (c) proprietary and confidential information, ideas, mediation, inventions, type, developmental or experimental work, modifications, techniques, designs, product plans, product integration plans, data, computer programs, other copyrightable works, know-how, and processes.

3. EXCEPTIONS TO CONFIDENTIAL INFORMATION
"Confidential Information" does not include information that: (a) was lawfully in the possession of you or any of your Representatives prior to its disclosure to you by the Company; (b) becomes available to you or your Representatives from a source other than the Company; (c) becomes generally available to the public through no fault of you or any of your Representatives in breach of this agreement; (d) was or is independently developed by you or your Representatives without use of, reference to, or based upon Confidential Information.

4. COMPELLED DISCLOSURE
Where the Recipient is ordered by a court of competent jurisdiction to do so or by another authority who has the power to order disclosure and who does so order disclosure or there is a statutory or other legal obligation to do so; provided that the Recipient shall (i) first notify the Disclosing Party in writing before any disclosure under such order or obligation is made except the Recipient shall not so notify the Disclosing Party where it is not permitted to do so as a matter of law (provided that it will so notify the Disclosing Party as soon as permissible thereafter); and (ii) ensure that the court of competent jurisdiction or such other authority is made aware, prior to the disclosure, of the confidential nature of the Confidential Information.

5. MISCELLANEOUS
This Agreement constitutes the entire agreement between the parties with respect to the subject matter hereof."""


def load_playbook_display_rules(path: str = PLAYBOOK_FINAL_FILE):
    """Load playbook for sidebar display."""
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, list) else []
    except Exception:
        return []


# ═══════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════

def main():
    st.title("AI Legal Assistant")
    st.caption(
        "Surgical-precision contract redlining powered by OpenAI  ·  "
        "Playbook-driven  ·  Minimalist modifications"
    )

    tabs = ["📄 Contract Review", "📚 Rule Learning", "🧠 Rule Management"]
    tab1, tab2, tab3 = st.tabs(tabs)

    with tab1:
        render_contract_review_tab()
    with tab2:
        render_rule_learning_tab()
    with tab3:
        render_rule_management_tab()


# ═══════════════════════════════════════════════════════════════════
# TAB 1: Contract Review (NEW unified pipeline)
# ═══════════════════════════════════════════════════════════════════

def render_contract_review_tab():
    if not UNIFIED_AVAILABLE:
        st.error(
            "Unified review pipeline not available. "
            "Please set the OPENAI_API_KEY environment variable and restart."
        )
        return

    # ── Sidebar: Playbook rules ──────────────────────────────────
    with st.sidebar:
        st.markdown("### 📋 Playbook Rules")
        all_playbooks = load_playbooks_for_display()

        # Per-playbook toggle
        if "selected_playbook_ids" not in st.session_state:
            st.session_state["selected_playbook_ids"] = [e["id"] for e in all_playbooks]

        select_col1, select_col2 = st.columns(2)
        with select_col1:
            if st.button("Select All", use_container_width=True, key="pb_all"):
                st.session_state["selected_playbook_ids"] = [e["id"] for e in all_playbooks]
                st.rerun()
        with select_col2:
            if st.button("Deselect All", use_container_width=True, key="pb_none"):
                st.session_state["selected_playbook_ids"] = []
                st.rerun()

        for entry in all_playbooks:
            eid = entry["id"]

            is_on = st.checkbox(
                entry.get("title", "Rule"),
                value=(eid in st.session_state["selected_playbook_ids"]),
                key=f"pb_cb_{eid}",
            )
            if is_on and eid not in st.session_state["selected_playbook_ids"]:
                st.session_state["selected_playbook_ids"].append(eid)
            elif not is_on and eid in st.session_state["selected_playbook_ids"]:
                st.session_state["selected_playbook_ids"].remove(eid)

            with st.expander(f"Details: {entry.get('title','')}", expanded=False):
                st.caption(
                    f"ID: {eid}  ·  "
                    f"Type: {entry.get('type', 'add_text')}  ·  "
                    f"File: `{entry.get('source_file', '')}`"
                )
                st.markdown(f"**Rule Summary:** {entry.get('summary', '')}")
                st.markdown("---")
                st.markdown(entry.get("markdown_body", ""))

        selected_count = len(st.session_state.get("selected_playbook_ids", []))
        st.caption(f"**{selected_count} / {len(all_playbooks)}** playbooks selected")

    # ── Main area ────────────────────────────────────────────────
    mode = "own_paper"
    generate_issues = True
    generate_word = True
    preset = "quality"
    os.environ["MODEL_PRESET"] = preset

    st.markdown("### 📄 Contract Text")
    uploaded = st.file_uploader(
        "Upload .docx or .txt (optional)",
        type=["docx", "txt"],
    )
    uploaded_text = parse_uploaded_contract(uploaded, store_source_docx=True) if uploaded else None
    if not uploaded:
        st.session_state.pop("review_source_docx_bytes", None)

    contract_text = st.text_area(
        "Contract Text",
        value=uploaded_text or build_default_contract(),
        height=400,
        help="Paste or edit the contract text here.",
    )

    st.divider()

    # ── Run button ───────────────────────────────────────────────
    run_col1, run_col2, run_col3 = st.columns([1, 2, 1])
    with run_col2:
        run = st.button("🚀 Run Review", type="primary", use_container_width=True)

    if not run:
        st.info(
            "Upload a contract or use the sample text above, "
            "then click **Run Review** to start."
        )
        return

    if not os.environ.get("OPENAI_API_KEY"):
        st.error("Missing OPENAI_API_KEY. Please set it in your shell before starting the app.")
        return

    if not contract_text.strip():
        st.error("Contract text is empty.")
        return

    endpoint = (
        os.environ.get("OPENAI_API_BASE")
        or os.environ.get("OPENAI_BASE_URL")
        or "https://api.openai.com/v1"
    )
    st.caption(f"LLM endpoint: `{endpoint}`")

    # ── Execute pipeline ─────────────────────────────────────────
    progress_bar = st.progress(0)
    status_area = st.empty()
    stages = {
        "parsing": 0.05,
        "analysis": 0.15,
        "analysis_done": 0.45,
        "execution": 0.55,
        "execution_done": 0.75,
        "issues": 0.80,
        "issues_done": 0.90,
        "done": 1.0,
    }

    def on_progress(stage, detail=""):
        pct = stages.get(stage, 0)
        progress_bar.progress(pct)
        status_area.markdown(f"**{detail}**" if detail else f"Stage: {stage}")

    try:
        selected_ids = st.session_state.get("selected_playbook_ids", [])
        if not selected_ids:
            st.error("No playbooks selected. Please enable at least one playbook in the sidebar.")
            return
        playbook_entries = load_playbooks_from_markdown(filter_ids=selected_ids)
        if not playbook_entries:
            st.error("No playbook rules found in playbooks/ directory")
            return
        on_progress("parsing", f"Loaded {len(playbook_entries)} of {len(selected_ids)} playbooks")

        result = unified_review_contract(
            contract_text=contract_text,
            playbook_entries=playbook_entries,
            mode=mode,
            generate_issues=generate_issues,
            progress_callback=on_progress,
            playbook_source="markdown",
        )
    except Exception as e:
        progress_bar.empty()
        status_area.empty()
        st.error(f"Review failed: {_friendly_review_error_message(e)}")
        import traceback
        with st.expander("Technical details", expanded=False):
            st.code(traceback.format_exc())
        return

    progress_bar.empty()
    status_area.empty()

    # ── Display results ──────────────────────────────────────────
    final_text = result.get("final_text", contract_text)
    modifications = result.get("modifications", [])
    analysis = result.get("analysis", [])
    issues = result.get("issues_list", [])
    summary = result.get("summary", {})
    llm_stats = result.get("llm_stats", {})
    executive_summary = result.get("executive_summary", "")
    compliance = result.get("compliance_score", {})
    step_trace = result.get("step_trace", [])

    # ── Summary cards ────────────────────────────────────────────
    st.markdown("---")
    st.markdown("## Review Results")

    if executive_summary:
        if mode == "counterparty":
            st.warning(f"⚠️ **Counterparty Paper Assessment:** {executive_summary}")
        else:
            st.success(f"✅ **Assessment:** {executive_summary}")

    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("Rules Checked", summary.get("total_rules_checked", len(analysis)))
    c2.metric("Compliant", summary.get("compliant", 0))
    c3.metric("Non-Compliant", summary.get("non_compliant", 0))
    c4.metric("Modifications", len(modifications))
    c5.metric("Risk", summary.get("overall_risk", "—").upper())

    if compliance:
        pct = compliance.get("percentage", 0)
        st.progress(pct / 100 if isinstance(pct, (int, float)) and pct <= 100 else 0)
        st.caption(f"Compliance: {pct}%  ({compliance.get('compliant',0)} / {compliance.get('total_rules',0)} rules)")

    # ── Tabs for details ─────────────────────────────────────────
    detail_tabs = st.tabs(["📝 Redline", "📊 Issues List", "🔍 Analysis Detail", "📈 LLM Stats", "🧠 Agent Thinking"])

    # ── Tab: Redline ─────────────────────────────────────────────
    with detail_tabs[0]:
        changed, ops_count = diff_changed(contract_text, final_text)
        if not changed:
            st.success("Contract is fully compliant — no modifications needed.")
        else:
            st.markdown(f"**{len(modifications)} modifications applied** ({ops_count} edit operations)")

            # Full redline diff
            diff_html = generate_diff_html(contract_text, final_text)
            st.markdown(
                f"<div class='clause-box' style='max-height:600px;overflow-y:auto;'>{diff_html}</div>",
                unsafe_allow_html=True,
            )

            # Per-modification detail with AI reasoning
            analysis_by_rule = {
                item.get("rule_id", ""): item for item in analysis
            }

            st.markdown("#### Modification Details")
            for i, mod in enumerate(modifications, 1):
                rule_id = mod.get("rule_id", "")
                sev = mod.get("severity", "P1")
                severity_color = {"P0": "RED", "P1": "YELLOW", "P2": "GREEN"}.get(sev, "YELLOW")
                icon = {"P0": "🔴", "P1": "🟡", "P2": "⚪"}.get(sev, "🟡")

                with st.expander(
                    f"{icon} {i}. {mod.get('rule_title', 'Modification')} — {severity_color} ({sev})",
                    expanded=True,
                ):
                    # Redline diff for this modification
                    orig = mod.get("original_fragment", "")
                    modified = mod.get("modified_fragment", "")
                    if orig and modified:
                        mini_diff = generate_diff_html(orig, modified)
                        st.markdown(f"<div class='clause-box'>{mini_diff}</div>", unsafe_allow_html=True)

                    st.markdown(f"**Explanation:** {mod.get('explanation', '—')}")

                    analysis_item = analysis_by_rule.get(rule_id)
                    rationale = analysis_item.get("rationale") or analysis_item.get("chain_of_thought") if analysis_item else None
                    if rationale:
                        st.markdown("**AI Rationale:**")
                        st.info(rationale)

                    # Show the modification plan
                    if analysis_item and analysis_item.get("modification_plan"):
                        with st.expander("Modification Plan (raw)", expanded=False):
                            st.json(analysis_item["modification_plan"])

        # Download buttons
        st.markdown("---")
        dl_col1, dl_col2, dl_col3 = st.columns(3)

        if generate_word and changed:
            with dl_col1:
                redline_buf = generate_redline_docx(
                    contract_text,
                    final_text,
                    issues_list=issues or None,
                    modifications=modifications or None,
                    title=None,
                    redline_heading=None,
                    include_issues_list=False,
                    source_docx_bytes=st.session_state.get("review_source_docx_bytes"),
                )
                st.download_button(
                    "📥 Download Redline (.docx)",
                    data=redline_buf,
                    file_name=f"redline_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            with dl_col2:
                clean_buf = generate_clean_docx(final_text, title="Contract — Clean Copy")
                st.download_button(
                    "📥 Download Clean Copy (.docx)",
                    data=clean_buf,
                    file_name=f"clean_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
        with dl_col3 if generate_word and changed else dl_col1:
            st.download_button(
                "📥 Download Modified Text (.txt)",
                data=final_text,
                file_name=f"modified_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
                mime="text/plain",
                use_container_width=True,
            )

    # ── Tab: Issues List ─────────────────────────────────────────
    with detail_tabs[1]:
        if not issues:
            st.info("No issues generated. Enable 'Generate Issues List' to see results.")
        else:
            st.markdown(f"### Issues List ({len(issues)} items)")

            sev_counts = {}
            for iss in issues:
                s = iss.get("severity", "P2")
                sev_counts[s] = sev_counts.get(s, 0) + 1
            ic1, ic2, ic3 = st.columns(3)
            ic1.metric("🔴 P0 Critical", sev_counts.get("P0", 0))
            ic2.metric("🟡 P1 Important", sev_counts.get("P1", 0))
            ic3.metric("⚪ P2 Minor", sev_counts.get("P2", 0))

            for iss in sorted(issues, key=lambda x: x.get("severity", "P2")):
                sev = iss.get("severity", "P2")
                css_class = {"P0": "issue-p0", "P1": "issue-p1", "P2": "issue-p2"}.get(sev, "issue-p2")
                icon = {"P0": "🔴", "P1": "🟡", "P2": "⚪"}.get(sev, "⚪")

                with st.expander(
                    f"{icon} [{sev}] {iss.get('title', iss.get('description','Issue'))}",
                    expanded=(sev == "P0"),
                ):
                    st.markdown(f"**Category:** {iss.get('category', '—')}")
                    st.markdown(f"**Clause:** {iss.get('clause_reference', '—')}")
                    st.markdown(f"**Description:** {iss.get('description', '—')}")
                    if iss.get("current_language"):
                        st.markdown(f"**Current Language:**")
                        st.code(iss["current_language"], language="text")
                    st.markdown(f"**Recommended Action:** {iss.get('recommended_action', '—')}")
                    st.caption(f"Status: {iss.get('status','—')}  ·  Rule: {iss.get('playbook_rule','—')}")

    # ── Tab: Analysis Detail ─────────────────────────────────────
    with detail_tabs[2]:
        st.markdown(f"### Rule-by-Rule Analysis ({len(analysis)} rules)")
        for item in analysis:
            status = item.get("compliance_status", "unknown")
            icon = {"compliant": "✅", "non_compliant": "🔴", "partially_compliant": "🟡"}.get(status, "❓")
            rule_title = item.get("rule_title", item.get("rule_id", "Rule"))
            source = item.get("source_playbook", "")

            with st.expander(f"{icon} {rule_title} — {status.upper()}", expanded=(status != "compliant")):
                st.caption(f"Rule ID: {item.get('rule_id','')}  ·  Type: {item.get('rule_type','')}  ·  Source: {source}")
                st.markdown(f"**Matched Clause:** {item.get('clause_location','—')}")
                if item.get("matched_clause"):
                    st.code(item["matched_clause"][:500], language="text")

                st.markdown("**Rationale:**")
                st.markdown(item.get("rationale") or item.get("chain_of_thought", "—"))

                if item.get("modification_needed"):
                    plan = item.get("modification_plan", {})
                    st.markdown("**Modification Plan:**")
                    st.json(plan)

    # ── Tab: LLM Stats ───────────────────────────────────────────
    with detail_tabs[3]:
        st.markdown("### LLM Call Statistics")
        total_calls = llm_stats.get("total_calls", 0)
        total_dur = llm_stats.get("total_duration", 0)
        total_tok = llm_stats.get("total_tokens", 0)

        sc1, sc2, sc3 = st.columns(3)
        sc1.metric("API Calls", total_calls)
        sc2.metric("Total Time", f"{total_dur:.1f}s")
        sc3.metric("Total Tokens", f"{total_tok:,}")

        for i, call in enumerate(llm_stats.get("calls", []), 1):
            st.markdown(
                f"**Call {i}:** `{call.get('task','')}` → "
                f"model=`{call.get('model','')}` · "
                f"{call.get('duration',0):.1f}s · "
                f"{call.get('tokens',0):,} tokens "
                f"(prompt={call.get('prompt_tokens',0):,} completion={call.get('completion_tokens',0):,})"
            )

    # ── Tab: Agent Thinking ──────────────────────────────────────
    with detail_tabs[4]:
        st.markdown("### Agent Step-by-Step Thinking")
        st.caption(
            "Structured summaries of each pipeline step "
            "(parsing → analysis → execution → issues)."
        )

        if not step_trace:
            st.info("No step trace found in this run.")
        else:
            for step_item in step_trace:
                step_title = f"{step_item.get('step', 'Step')} · {step_item.get('name', '')}".strip()
                engine = step_item.get("engine", "Unknown engine")
                with st.expander(step_title, expanded=True):
                    st.markdown(f"**Engine:** `{engine}`")

                    thinking = step_item.get("thinking", [])
                    if thinking:
                        st.markdown("**Reasoning Summary:**")
                        for line in thinking:
                            st.markdown(f"- {line}")

                    output_obj = step_item.get("output")
                    if output_obj:
                        st.markdown("**Structured Output:**")
                        st.json(output_obj)


# ═══════════════════════════════════════════════════════════════════
# TAB 2: Rule Learning (OpenAI provider)
# ═══════════════════════════════════════════════════════════════════

def render_rule_learning_tab():
    st.markdown("### 📚 Learn Modification Rules from Historical Contracts")
    st.caption(
        "Upload 'before' and 'after' contract pairs — "
        "AI will automatically analyze differences and extract reusable rules."
    )

    if not LEARNING_AVAILABLE:
        st.warning(
            "Rule learning requires OpenAI provider. "
            "Please set OPENAI_API_KEY and ensure providers/rule_learning.py is available."
        )
        return

    st.divider()
    col1, col2 = st.columns(2)

    with col1:
        st.markdown("#### 📄 Original Contract (Before)")
        before_source = st.radio(
            "Input method", ["Paste text", "Upload file"],
            horizontal=True, key="before_src",
        )
        if before_source == "Upload file":
            bf = st.file_uploader("Upload original (.docx/.txt)", type=["docx", "txt"], key="bf_up")
            before_text = parse_uploaded_contract(bf) if bf else ""
            if before_text:
                st.text_area("Preview", value=before_text, height=250, disabled=True, key="bf_prev")
        else:
            before_text = st.text_area("Paste original text", height=250, key="bf_paste")

    with col2:
        st.markdown("#### ✏️ Modified Contract (After)")
        after_source = st.radio(
            "Input method", ["Paste text", "Upload file"],
            horizontal=True, key="after_src",
        )
        if after_source == "Upload file":
            af = st.file_uploader("Upload modified (.docx/.txt)", type=["docx", "txt"], key="af_up")
            after_text = parse_uploaded_contract(af) if af else ""
            if after_text:
                st.text_area("Preview", value=after_text, height=250, disabled=True, key="af_prev")
        else:
            after_text = st.text_area("Paste modified text", height=250, key="af_paste")

    st.divider()
    ci1, ci2 = st.columns([2, 1])
    with ci1:
        case_name = st.text_input("Case Name", value="NDA Modification Case", key="case_name")
    with ci2:
        auto_save = st.checkbox("Auto-save rules", value=True, key="auto_save")

    if st.button("🔍 Analyse & Extract Rules", type="primary", use_container_width=True):
        if not before_text or not after_text:
            st.error("Provide both before and after texts.")
            return

        with st.status("Analysing differences…", expanded=True) as status:
            st.write("Comparing texts…")
            try:
                result = learn_from_contract_diff(
                    before_text=before_text,
                    after_text=after_text,
                    case_name=case_name,
                    auto_save=auto_save,
                )
                extracted = result["extracted_rules"]
                status.update(label=f"Extracted {len(extracted)} rules", state="complete", expanded=False)
            except Exception as e:
                status.update(label="Extraction failed", state="error")
                st.error(str(e))
                return

        if extracted:
            st.success(f"Extracted {len(extracted)} reusable rules!")
            for i, rule in enumerate(extracted, 1):
                with st.expander(f"Rule {i}: {rule.name}", expanded=True):
                    rc1, rc2 = st.columns(2)
                    with rc1:
                        st.markdown(f"**Type:** `{rule.type}`")
                        st.markdown(f"**Trigger:** {rule.trigger}")
                        st.markdown(f"**Action:** {rule.action}")
                        st.markdown(f"**Confidence:** {rule.confidence:.0%}")
                    with rc2:
                        st.code(rule.exact_wording, language="text")
                        st.markdown(f"**Rationale:** {rule.rationale}")


# ═══════════════════════════════════════════════════════════════════
# TAB 3: Rule Management (OpenAI provider)
# ═══════════════════════════════════════════════════════════════════

def render_rule_management_tab():
    st.markdown("### 🧠 Learned Rules Management")

    if not LEARNING_AVAILABLE:
        st.warning("Rule management requires OpenAI provider.")
        return

    try:
        stats = get_learned_rules_stats()
        store = load_learned_rules()
    except Exception as e:
        st.error(f"Failed to load rules: {e}")
        return

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Total Rules", stats["total_rules"])
    c2.metric("Enabled", stats["enabled_rules"])
    c3.metric("Applications", stats["total_applications"])
    c4.metric("Cases Learned", stats["learning_statistics"].get("total_cases_learned", 0))

    st.divider()

    if not store.rules:
        st.info("No learned rules yet. Use the Rule Learning tab to get started.")
        return

    type_filter = st.multiselect(
        "Filter by type",
        ["add_text", "replace_text", "checklist", "conditional"],
        key="mgmt_filter",
    )

    for rule in store.rules:
        if type_filter and rule.type not in type_filter:
            continue

        sicon = "✅" if rule.enabled else "⏸️"
        cicon = "🟢" if rule.confidence >= 0.8 else ("🟡" if rule.confidence >= 0.5 else "🔴")

        with st.expander(f"{sicon} {rule.name}  |  {cicon} {rule.confidence:.0%}  |  {rule.source_case_name}"):
            dc1, dc2 = st.columns([2, 1])
            with dc1:
                st.markdown(f"**ID:** `{rule.id}`  ·  **Type:** `{rule.type}`")
                st.markdown(f"**Trigger:** {rule.trigger}")
                st.markdown(f"**Action:** {rule.action}")
                st.code(rule.exact_wording, language="text")
            with dc2:
                st.markdown(f"**Confidence:** {rule.confidence:.0%}")
                st.markdown(f"**Applied:** {rule.times_applied} times")
                st.markdown(f"**Learned:** {rule.learned_at[:10]}")
                if rule.enabled:
                    if st.button("⏸️ Disable", key=f"dis_{rule.id}"):
                        rule.enabled = False
                        save_learned_rules(store)
                        st.rerun()
                else:
                    if st.button("▶️ Enable", key=f"en_{rule.id}"):
                        rule.enabled = True
                        save_learned_rules(store)
                        st.rerun()


# ═══════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    main()
